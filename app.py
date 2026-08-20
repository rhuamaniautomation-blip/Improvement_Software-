#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
================================================================================
SISTEMA DE GESTION DOCUMENTAL - MoC | Mejora A3 | Simple Kaizen
Version 10.0.0 - Generación Nativa sin Dependencia de Templates
================================================================================
Diseñado por: CAVA - Especialistas en Robotica y Automatizacion
Desarrollador: Roger Huamani
Version: 10.0.0
Fecha: Agosto 2026
================================================================================
"""
import streamlit as st
import os
import json
import re
import uuid
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor as DocxRGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak, KeepTogether
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY, TA_RIGHT
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# =============================================================================
# CONFIGURACION INICIAL
# =============================================================================
st.set_page_config(
    page_title="Gestión Documental - MoC | A3 | Kaizen",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

BASE_DIR = Path(__file__).parent if "__file__" in dir() else Path(".")
DATA_DIR = BASE_DIR / "data"
HISTORY_FILE = DATA_DIR / "history.json"
CONFIG_FILE = DATA_DIR / "config.json"
DATA_DIR.mkdir(exist_ok=True)

# =============================================================================
# CSS PERSONALIZADO
# =============================================================================
CUSTOM_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
html, body, [class*="css"] { font-family: 'Inter', 'Segoe UI', sans-serif !important; }
.main-header {
    background: linear-gradient(135deg, #1a5f7a 0%, #2e8bc0 100%);
    padding: 2rem; border-radius: 12px; color: white;
    margin-bottom: 2rem; box-shadow: 0 4px 20px rgba(26, 95, 122, 0.3);
}
.main-header h1 { color: white !important; font-weight: 700 !important; margin-bottom: 0.5rem !important; }
.main-header p { color: rgba(255,255,255,0.9) !important; font-size: 1.1rem !important; }
.doc-card {
    background: white; border-radius: 16px; padding: 2rem;
    border: 2px solid #e2e8f0; transition: all 0.3s ease;
    cursor: pointer; height: 100%;
}
.doc-card:hover {
    border-color: #1a5f7a; transform: translateY(-4px);
    box-shadow: 0 12px 40px rgba(26, 95, 122, 0.15);
}
.doc-card-moc { border-left: 5px solid #1a5f7a; }
.doc-card-a3 { border-left: 5px solid #10b981; }
.doc-card-kaizen { border-left: 5px solid #f59e0b; }
.stButton > button {
    border-radius: 10px !important; font-weight: 600 !important;
    padding: 0.75rem 2rem !important; transition: all 0.2s ease !important;
}
.stButton > button:hover {
    transform: translateY(-2px); box-shadow: 0 4px 12px rgba(0,0,0,0.15);
}
.section-header {
    background: #f1f5f9; padding: 1rem 1.5rem;
    border-radius: 10px; margin: 1.5rem 0 1rem 0;
    border-left: 4px solid #1a5f7a;
}
.section-header h3 { margin: 0 !important; color: #1e293b !important; font-weight: 600 !important; }
.gemini-badge {
    display: inline-block; background: #e0e7ff; color: #4338ca;
    padding: 0.25rem 0.75rem; border-radius: 20px;
    font-size: 12px; font-weight: 600; margin-left: 0.5rem;
}
.history-item {
    background: white; border: 1px solid #e2e8f0;
    border-radius: 10px; padding: 1rem; margin: 0.5rem 0;
}
.app-footer {
    text-align: center; padding: 2rem; margin-top: 3rem;
    border-top: 1px solid #e2e8f0; color: #64748b;
}
.auto-correct-badge {
    display: inline-flex; align-items: center;
    background: #dcfce7; color: #166534;
    padding: 0.25rem 0.75rem; border-radius: 20px;
    font-size: 11px; font-weight: 600; margin-bottom: 0.5rem;
}
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
[data-testid="stSidebar"] { background: #1e293b !important; }
[data-testid="stSidebar"] .stMarkdown { color: #94a3b8 !important; }
[data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 { color: #f8fafc !important; }
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

# =============================================================================
# PERSISTENCIA LOCAL
# =============================================================================
class LocalStorage:
    @staticmethod
    def save_config(config):
        try:
            with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            st.error(f"Error guardando configuración: {e}")
            return False

    @staticmethod
    def load_config():
        try:
            if CONFIG_FILE.exists():
                with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            st.warning(f"Error cargando configuración: {e}")
        return None

    @staticmethod
    def save_history(history):
        try:
            with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
                json.dump(history, f, ensure_ascii=False, indent=2)
            return True
        except Exception as e:
            st.error(f"Error guardando historial: {e}")
            return False

    @staticmethod
    def load_history():
        try:
            if HISTORY_FILE.exists():
                with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except Exception as e:
            st.warning(f"Error cargando historial: {e}")
        return {"documents": []}

# =============================================================================
# UTILIDADES
# =============================================================================
class Utils:
    @staticmethod
    def format_date():
        meses = {1:"ENERO",2:"FEBRERO",3:"MARZO",4:"ABRIL",5:"MAYO",6:"JUNIO",
                 7:"JULIO",8:"AGOSTO",9:"SEPTIEMBRE",10:"OCTUBRE",11:"NOVIEMBRE",12:"DICIEMBRE"}
        now = datetime.now()
        return f"{meses[now.month]} {now.year}"

    @staticmethod
    def format_date_short():
        now = datetime.now()
        return now.strftime("%d/%m/%Y")

    @staticmethod
    def generate_doc_number(doc_type):
        config = st.session_state.config
        key = f"last_{doc_type}_number"
        config[key] = config.get(key, 0) + 1
        st.session_state.config = config
        LocalStorage.save_config(config)
        now = datetime.now()
        prefix = {"moc": "MOC", "a3": "A3", "kaizen": "KZN"}
        return f"{prefix.get(doc_type, 'DOC')}-{now.year}{now.month:02d}{now.day:02d}-{config[key]:04d}"

    @staticmethod
    def add_to_history(doc_info):
        history = st.session_state.history
        doc_info["id"] = str(uuid.uuid4())
        doc_info["timestamp"] = datetime.now().isoformat()
        history["documents"].insert(0, doc_info)
        st.session_state.history = history
        LocalStorage.save_history(history)

    @staticmethod
    def delete_from_history(doc_id):
        history = st.session_state.history
        history["documents"] = [d for d in history["documents"] if d.get("id") != doc_id]
        st.session_state.history = history
        LocalStorage.save_history(history)

    @staticmethod
    def correct_spelling_basic(text):
        if not text or not text.strip():
            return text
        corrections = {
            "tecnico": "técnico", "Tecnico": "Técnico",
            "tecnologia": "tecnología", "Tecnologia": "Tecnología",
            "produccion": "producción", "Produccion": "Producción",
            "implementacion": "implementación", "Implementacion": "Implementación",
            "evaluacion": "evaluación", "Evaluacion": "Evaluación",
            "operacion": "operación", "Operacion": "Operación",
            "condicion": "condición", "Condicion": "Condición",
            "modificacion": "modificación", "Modificacion": "Modificación",
            "verificacion": "verificación", "Verificacion": "Verificación",
            "capacitacion": "capacitación", "Capacitacion": "Capacitación",
            "documentacion": "documentación", "Documentacion": "Documentación",
            "estandarizacion": "estandarización", "Estandarizacion": "Estandarización",
            "optimizacion": "optimización", "Optimizacion": "Optimización",
            "identificacion": "identificación", "Identificacion": "Identificación",
            "clasificacion": "clasificación", "Clasificacion": "Clasificación",
            "notificacion": "notificación", "Notificacion": "Notificación",
            "coordinacion": "coordinación", "Coordinacion": "Coordinación",
            "aprobacion": "aprobación", "Aprobacion": "Aprobación",
            "revison": "revisión", "Revison": "Revisión",
            "ejecucion": "ejecución", "Ejecucion": "Ejecución",
            "inspeccion": "inspección", "Inspeccion": "Inspección",
            "proteccion": "protección", "Proteccion": "Protección",
            "deteccion": "detección", "Deteccion": "Detección",
            "prevencion": "prevención", "Prevencion": "Prevención",
            "intervencion": "intervención", "Intervencion": "Intervención",
            "supervision": "supervisión", "Supervision": "Supervisión",
            "comunicacion": "comunicación", "Comunicacion": "Comunicación",
            "organizacion": "organización", "Organizacion": "Organización",
            "planificacion": "planificación", "Planificacion": "Planificación",
            "calificacion": "calificación", "Calificacion": "Calificación",
            "certificacion": "certificación", "Certificacion": "Certificación",
            "validacion": "validación", "Validacion": "Validación",
            "calibracion": "calibración", "Calibracion": "Calibración",
            "configuracion": "configuración", "Configuracion": "Configuración",
            "programacion": "programación", "Programacion": "Programación",
            "automatizacion": "automatización", "Automatizacion": "Automatización",
            "integracion": "integración", "Integracion": "Integración",
            "funcion": "función", "Funcion": "Función",
            "relacion": "relación", "Relacion": "Relación",
            "conexion": "conexión", "Conexion": "Conexión",
            "direccion": "dirección", "Direccion": "Dirección",
            "seleccion": "selección", "Seleccion": "Selección",
            "distribucion": "distribución", "Distribucion": "Distribución",
            "construccion": "construcción", "Construccion": "Construcción",
            "instruccion": "instrucción", "Instruccion": "Instrucción",
            "reduccion": "reducción", "Reduccion": "Reducción",
            "traduccion": "traducción", "Traduccion": "Traducción",
            "maquina": "máquina", "Maquina": "Máquina",
            "maquinas": "máquinas", "Maquinas": "Máquinas",
            "podria": "podría", "Podria": "Podría",
            "habria": "habría", "Habria": "Habría",
            "seria": "sería", "Seria": "Sería",
            "tendria": "tendría", "Tendria": "Tendría",
            "haria": "haría", "Haria": "Haría",
            "daria": "daría", "Daria": "Daría",
            "estaria": "estaría", "Estaria": "Estaría",
            "deberia": "debería", "Deberia": "Debería",
            "mas": "más", "Mas": "Más",
            "aun": "aún", "Aun": "Aún",
            "tambien": "también", "Tambien": "También",
            "asi": "así", "Asi": "Así",
            "aqui": "aquí", "Aqui": "Aquí",
            "alli": "allí", "Alli": "Allí",
            "despues": "después", "Despues": "Después",
            "ademas": "además", "Ademas": "Además",
            "segun": "según", "Segun": "Según",
            "numero": "número", "Numero": "Número",
            "maximo": "máximo", "Maximo": "Máximo",
            "minimo": "mínimo", "Minimo": "Mínimo",
            "optimo": "óptimo", "Optimo": "Óptimo",
            "ultimo": "último", "Ultimo": "Último",
            "periodo": "período", "Periodo": "Período",
            "area": "área", "Area": "Área",
            "dia": "día", "Dia": "Día",
            "manana": "mañana", "Manana": "Mañana",
            "proximo": "próximo", "Proximo": "Próximo",
            "analisis": "análisis", "Analisis": "Análisis",
            "metodo": "método", "Metodo": "Método",
            "parametro": "parámetro", "Parametro": "Parámetro",
            "parametros": "parámetros", "Parametros": "Parámetros",
            "caracteristica": "característica", "Caracteristica": "Característica",
            "caracteristicas": "características", "Caracteristicas": "Características",
            "especifico": "específico", "Especifico": "Específico",
            "especifica": "específica", "Especifica": "Específica",
            "generico": "genérico", "Generico": "Genérico",
            "electronico": "electrónico", "Electronico": "Electrónico",
            "electrico": "eléctrico", "Electrico": "Eléctrico",
            "hidraulico": "hidráulico", "Hidraulico": "Hidráulico",
            "neumatico": "neumático", "Neumatico": "Neumático",
            "termico": "térmico", "Termico": "Térmico",
            "quimico": "químico", "Quimico": "Químico",
            "fisico": "físico", "Fisico": "Físico",
            "biologico": "biológico", "Biologico": "Biológico",
            "version": "versión", "Version": "Versión",
            "descripcion": "descripción", "Descripcion": "Descripción",
            "solucion": "solución", "Solucion": "Solución",
            "situacion": "situación", "Situacion": "Situación",
            "presentacion": "presentación", "Presentacion": "Presentación",
            "revision": "revisión", "Revision": "Revisión",
            "habilitacion": "habilitación", "Habilitacion": "Habilitación",
            "limites": "límites", "Limites": "Límites",
            "limite": "límite", "Limite": "Límite",
            "linea": "línea", "Linea": "Línea",
            "unico": "único", "Unico": "Único",
            "facil": "fácil", "Facil": "Fácil",
            "dificil": "difícil", "Dificil": "Difícil",
            "rapido": "rápido", "Rapido": "Rápido",
            "trabagar": "trabajar",
            "podra": "podrá",
            "esta": "está",
        }
        result = text
        for wrong, correct in corrections.items():
            result = result.replace(wrong, correct)
        result = re.sub(r'  +', ' ', result)
        result = re.sub(r' ([.,;:!?])', r'\1', result)
        return result

# =============================================================================
# SERVICIO GEMINI API
# =============================================================================
class GeminiService:
    MODELS = {
        "gemini-2.5-flash": {"name": "Gemini 2.5 Flash", "desc": "Rápido y recomendado"},
        "gemini-2.5-pro": {"name": "Gemini 2.5 Pro", "desc": "Máxima calidad"},
    }
    DEPRECATED_MODELS = {
        "gemini-1.5-flash-lite", "gemini-1.5-flash", "gemini-1.5-pro",
        "gemini-1.0-pro", "gemini-2.0-flash", "gemini-2.0-flash-lite"
    }
    DEFAULT_MODEL = "gemini-2.5-flash"

    def __init__(self, api_key="", model=None):
        self.api_key = api_key
        if model is None or model in self.DEPRECATED_MODELS or model not in self.MODELS:
            self.model = self.DEFAULT_MODEL
        else:
            self.model = model
        self.base_url = "https://generativelanguage.googleapis.com/v1beta/models"

    def _call_api(self, prompt, temperature=0.3, max_tokens=4096):
        if not self.api_key:
            raise ValueError("API Key no configurada")
        import requests
        url = f"{self.base_url}/{self.model}:generateContent?key={self.api_key}"
        payload = {
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {"temperature": temperature, "maxOutputTokens": max_tokens}
        }
        try:
            response = requests.post(url, json=payload, timeout=180)
            response.raise_for_status()
            result = response.json()
            if "candidates" in result and len(result["candidates"]) > 0:
                return result["candidates"][0]["content"]["parts"][0]["text"]
            return ""
        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 404:
                raise Exception(f"Error 404: Modelo '{self.model}' no disponible. Verifique que la 'Generative Language API' esté habilitada.")
            elif e.response.status_code == 403:
                raise Exception("Error 403: Acceso denegado. Verifique API Key.")
            else:
                raise Exception(f"Error HTTP {e.response.status_code}: {e.response.text}")

    def _extract_json(self, text):
        json_match = re.search(r'```json\s*(.*?)\s*```', text, re.DOTALL)
        if json_match:
            try:
                return json.loads(json_match.group(1))
            except json.JSONDecodeError:
                pass
        json_match = re.search(r'\{.*\}', text, re.DOTALL)
        if json_match:
            try:
                return json.loads(json_match.group())
            except json.JSONDecodeError:
                pass
        return {"generated_text": text, "error": "No se pudo extraer JSON válido"}

    def generate_moc(self, problem, context="", equipo="", alternativas=""):
        if not self.api_key:
            st.error("❌ API Key no configurada")
            return None
        prompt = f"""Actúa como un Ingeniero Senior de Ingeniería, Automatización, Mantenimiento, Seguridad de Procesos y Gestión del Cambio (MoC) en una planta industrial de manufactura. Tienes 20 años de experiencia.

CONTEXTO DEL PROBLEMA IDENTIFICADO:
{problem}

INFORMACIÓN ADICIONAL:
{context}

EQUIPO INVOLUCRADO:
{equipo}

ALTERNATIVAS CONSIDERADAS (si existen):
{alternativas if alternativas else 'No proporcionadas.'}

INSTRUCCIONES CRÍTICAS:
1. Redacta SIEMPRE en español técnico y corporativo.
2. Usa enfoque de ingeniería industrial, automatización, calidad, seguridad y confiabilidad.
3. NO pidas información adicional. Si falta algún dato, asume la mejor alternativa técnicamente razonable.
4. Redacta en párrafos bien estructurados de 4-6 oraciones con conectores lógicos.
5. Evita listas, excepto en "alternativas_consideradas".
6. Considera riesgos operativos, de calidad, seguridad, productividad, mantenimiento y cumplimiento normativo.
7. Cuando aplique, considera PLC, HMI, SCADA, sensores, instrumentación, ISO 13849, enclavamientos.
8. Si el cambio es de software/automatización, indica que no modifica la integridad mecánica.
9. ORTOGRAFÍA IMPECABLE: Tildes correctas en todas las palabras.
10. PROHIBIDO usar frases genéricas. TODO el contenido debe basarse EXCLUSIVAMENTE en el problema del usuario.

Genera en ESPAÑOL formato JSON con esta estructura EXACTA:
{{
  "moc_title": "Título técnico conciso (máx. 12 palabras)",
  "condicion_actual": "SLIDE 3. Describe detalladamente cómo opera actualmente el sistema. Mínimo 200 palabras.",
  "condicion_propuesta": "SLIDE 3. Describe detalladamente la solución propuesta. Mínimo 200 palabras.",
  "justificacion_moc": "SLIDE 4. Explica la oportunidad de mejora que motiva el cambio. Mínimo 150 palabras.",
  "descripcion_problema": "SLIDE 5. Desarrolla técnicamente el problema con causas, efectos y consecuencias. Mínimo 250 palabras.",
  "razones_cambio": "SLIDE 6. Explica las razones técnicas que justifican la modificación. Mínimo 200 palabras.",
  "alternativas_consideradas": "SLIDE 4. Genera TRES alternativas con viñetas '❖': Alternativa 1 (mantener actual), Alternativa 2 (solución parcial), Alternativa 3 (seleccionada).",
  "plan_retorno": "SLIDE 4. Párrafo único de 3 líneas indicando cómo retornar a condición original en caso de falla.",
  "recursos": "SLIDE 8. Describe recursos necesarios: Ingeniería, Automatización, Mantenimiento, Calidad, Producción, Seguridad, Proveedor, Materiales, Software, Validación. Mínimo 150 palabras.",
  "plan_implementacion": "SLIDE 8. Párrafo describiendo etapas: evaluación, ingeniería, programación, instalación, pruebas, validación y liberación. Mínimo 150 palabras.",
  "tiempo_duracion": "SLIDE 8. Estima duración total considerando: Ingeniería, Aprobaciones, Compras, Soporte proveedor, Instalación, Programación, Validación.",
  "riesgos_controles": [{{"riesgo": "...", "control": "...", "plazo": "..."}}],
  "riesgos_shes": [{{"riesgo": "...", "control": "...", "plazo": "..."}}],
  "impacto_esperado": "Resumen ejecutivo de 1-2 párrafos con beneficios en Seguridad, Calidad, Productividad, Confiabilidad, Mantenimiento, Costos, Cumplimiento normativo.",
  "resumen_ejecutivo": "Párrafo corto dirigido a aprobadores explicando por qué aprobar el cambio.",
  "checklist_360": [
    {{"numero": 1, "factor": "Interacción o impacto con otras áreas/procesos", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 2, "factor": "Cambios en los procedimientos operativos, arranque y parada", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 3, "factor": "Parámetros operativos y límites de control", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 4, "factor": "Cambios en interfaces hombre-máquina y gestión de alarmas", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 5, "factor": "Compatibilidad de materiales, sustancias y equipos", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 6, "factor": "Exposición ocupacional (ruido, polvo, ergonomía, etc.)", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 7, "factor": "Requerimientos de EPP y su compatibilidad", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 8, "factor": "Escenarios de emergencia y capacidad de respuesta", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 9, "factor": "Impacto en el almacenamiento y tránsito interno/externo", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 10, "factor": "Impactos ambientales y generación de residuos", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 11, "factor": "Impacto en la calidad del producto o servicio", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 12, "factor": "Cambios en roles, competencias y carga de trabajo", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 13, "factor": "Integridad de equipos, protecciones y sistemas de control", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 14, "factor": "Cumplimiento legal, normativo y permisos aplicables", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 15, "factor": "Cambios en las condiciones para trabajos especiales", "aplica": "SI/NO", "descripcion": "..."}},
    {{"numero": 16, "factor": "Cambios sucesivos que incrementan el riesgo global", "aplica": "SI/NO", "descripcion": "..."}}
  ],
  "documentos_impactados": [
    {{"numero": 1, "documento": "JSERA - IPERC", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 2, "documento": "Procedimiento de Trabajo, Instructivo/PO", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 3, "documento": "Formato/Checklist operativos", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 4, "documento": "Matriz de EPP", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 5, "documento": "MSDS de sustancias involucradas", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 6, "documento": "Mapa de Riesgos", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 7, "documento": "Plan de emergencias", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 8, "documento": "Plan de Mantenimiento", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 9, "documento": "Matriz de impactos ambientales", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 10, "documento": "Plan Monitoreos SSO requeridos", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 11, "documento": "Plan de tráfico", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 12, "documento": "Matriz de competencias, plan de entrenamiento", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 13, "documento": "Plan de calidad", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 14, "documento": "Planos y diagramas (layout, P&ID)", "aplica": "SI/NO", "modificacion": "..."}},
    {{"numero": 15, "documento": "Licencias y permisos aplicables", "aplica": "SI/NO", "modificacion": "..."}}
  ]
}}

Responde SOLO con JSON válido."""
        try:
            text = self._call_api(prompt, temperature=0.4, max_tokens=16000)
            result = self._extract_json(text)
            for key in result:
                if isinstance(result[key], str):
                    result[key] = Utils.correct_spelling_basic(result[key])
                elif isinstance(result[key], list):
                    for item in result[key]:
                        if isinstance(item, dict):
                            for k in item:
                                if isinstance(item[k], str):
                                    item[k] = Utils.correct_spelling_basic(item[k])
            return result
        except Exception as e:
            st.error(f"❌ Error API: {e}")
            return None

    def generate_a3(self, problem, context=""):
        if not self.api_key:
            st.error("❌ API Key no configurada")
            return None
        prompt = f"""Eres un experto senior en metodología A3 Lean con 15 años de experiencia. Redactas documentos A3 con redacción humanizada, técnica y profesional.
INSTRUCCIONES: Usa EXCLUSIVAMENTE el contexto del problema. NO inventes problemas genéricos. Ortografía impecable.
PROBLEMA: {problem}
CONTEXTO: {context}
Genera JSON con: titulo, antecedentes, problema_actual, analisis_situacion, objetivos, analisis_causa_raiz, contramedidas, resultados_esperados, plan_seguimiento, lecciones_aprendidas, estandarizacion."""
        try:
            text = self._call_api(prompt, temperature=0.4, max_tokens=8192)
            result = self._extract_json(text)
            for key in result:
                if isinstance(result[key], str):
                    result[key] = Utils.correct_spelling_basic(result[key])
            return result
        except Exception as e:
            st.error(f"❌ Error API: {e}")
            return None

    def generate_kaizen(self, activity, context=""):
        if not self.api_key:
            st.error("❌ API Key no configurada")
            return None
        prompt = f"""Eres un experto en Kaizen y Lean Manufacturing. Redactas registros Kaizen con redacción humanizada y práctica.
INSTRUCCIONES: Usa EXCLUSIVAMENTE el contexto. NO inventes problemas. Ortografía impecable.
ACTIVIDAD: {activity}
CONTEXTO: {context}
Genera JSON con: titulo, area, descripcion_problema, solucion, beneficios, tipo_desperdicio, impacto_bto, proximos_pasos, leader, team_members."""
        try:
            text = self._call_api(prompt, temperature=0.4, max_tokens=4096)
            result = self._extract_json(text)
            for key in result:
                if isinstance(result[key], str):
                    result[key] = Utils.correct_spelling_basic(result[key])
            return result
        except Exception as e:
            st.error(f"❌ Error API: {e}")
            return None

    def translate_document(self, data):
        if not self.api_key:
            return data
        prompt = f"""Traduce del español al inglés profesional industrial: {json.dumps(data, ensure_ascii=False, indent=2)}. Responde SOLO el JSON traducido."""
        try:
            text = self._call_api(prompt, temperature=0.2, max_tokens=8192)
            return self._extract_json(text)
        except:
            return data

    def correct_spelling(self, text):
        if not self.api_key or not text.strip():
            return Utils.correct_spelling_basic(text)
        prompt = f"""Corrige ortografía, gramática y puntuación. Mantén el significado técnico. Asegura tildes correctas. Devuelve SOLO el texto corregido.
TEXTO:
{text}"""
        try:
            corrected = self._call_api(prompt, temperature=0.2, max_tokens=4096).strip()
            return Utils.correct_spelling_basic(corrected)
        except:
            return Utils.correct_spelling_basic(text)

# =============================================================================
# GENERADOR NATIVO DE DOCUMENTOS (SIN TEMPLATES)
# =============================================================================
class NativeDocumentGenerator:
    """Genera documentos Word nativos sin depender de templates externos"""
    
    @staticmethod
    def _set_cell_shading(cell, color):
        """Aplica color de fondo a una celda"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    @staticmethod
    def _add_header_row(table, row_idx, texts, bg_color="1A5F7A"):
        """Aplica formato de encabezado a una fila"""
        for i, text in enumerate(texts):
            if i < len(table.rows[row_idx].cells):
                cell = table.rows[row_idx].cells[i]
                cell.text = text
                NativeDocumentGenerator._set_cell_shading(cell, bg_color)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = DocxRGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(10)
    
    @staticmethod
    def _add_page_header(doc, title, subtitle=""):
        """Agrega encabezado estándar de página"""
        # Título principal
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = DocxRGBColor(0x1A, 0x5F, 0x7A)
            run.font.size = Pt(22)
        
        if subtitle:
            sub = doc.add_paragraph(subtitle)
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in sub.runs:
                run.font.color.rgb = DocxRGBColor(0x64, 0x74, 0x8B)
                run.font.size = Pt(11)
        
        doc.add_paragraph()
    
    @staticmethod
    def _add_section_title(doc, title, level=1):
        """Agrega título de sección con formato"""
        heading = doc.add_heading(title, level=level)
        for run in heading.runs:
            run.font.color.rgb = DocxRGBColor(0x1A, 0x5F, 0x7A)
        return heading
    
    @staticmethod
    def _add_content_paragraph(doc, text, bold=False, italic=False):
        """Agrega párrafo de contenido"""
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        run.bold = bold
        run.italic = italic
        p.paragraph_format.space_after = Pt(6)
        return p
    
    @staticmethod
    def generate_moc_docx(data, images=None):
        """Genera documento MoC en formato A4 Word nativo"""
        doc = Document()
        
        # Configurar márgenes A4
        for section in doc.sections:
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # ===== PORTADA =====
        NativeDocumentGenerator._add_page_header(
            doc, 
            "MANAGEMENT OF CHANGE (MoC)",
            "Gestión del Cambio"
        )
        
        # Tabla de información general
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        info_fields = [
            ("Título de la MoC:", data.get('moc_title', '')),
            ("Número:", data.get('moc_number', '')),
            ("Fecha:", data.get('fecha', '')),
            ("Naturaleza:", data.get('naturaleza', '').upper()),
            ("Originador:", data.get('originador', '')),
        ]
        for i, (label, value) in enumerate(info_fields):
            cell_label = table.rows[i].cells[0]
            cell_value = table.rows[i].cells[1]
            cell_label.text = label
            cell_value.text = value
            NativeDocumentGenerator._set_cell_shading(cell_label, "E2E8F0")
            for p in cell_label.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            for p in cell_value.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # Equipo de revisión
        NativeDocumentGenerator._add_section_title(doc, "Equipo de Revisión", level=2)
        equipo_text = f"""
• Producción: {data.get('produccion', 'N/A')}
• Specialist SHES: {data.get('specialist_shes', 'N/A')}
• Mantenimiento: {data.get('mantenimiento', 'N/A')}
• Revisores: {data.get('revisores', 'N/A')}
• Experto Aprobador: {data.get('experto_aprobador', 'N/A')}
"""
        NativeDocumentGenerator._add_content_paragraph(doc, equipo_text.strip())
        
        doc.add_page_break()
        
        # ===== CONDICIÓN ACTUAL vs PROPUESTA =====
        NativeDocumentGenerator._add_section_title(doc, "1. Condición Actual vs Condición Propuesta", level=1)
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        # Encabezados
        NativeDocumentGenerator._add_header_row(table, 0, ["CONDICIÓN ACTUAL", "CONDICIÓN PROPUESTA"])
        
        # Contenido
        table.rows[1].cells[0].text = data.get('condicion_actual', '')
        table.rows[1].cells[1].text = data.get('condicion_propuesta', '')
        for cell in table.rows[1].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # ===== JUSTIFICACIÓN =====
        NativeDocumentGenerator._add_section_title(doc, "2. Justificación de la MoC", level=1)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('justificacion_moc', ''))
        
        # ===== DESCRIPCIÓN DEL PROBLEMA =====
        NativeDocumentGenerator._add_section_title(doc, "3. Descripción del Problema", level=1)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('descripcion_problema', ''))
        
        # ===== RAZONES DEL CAMBIO =====
        NativeDocumentGenerator._add_section_title(doc, "4. Razones del Cambio", level=1)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('razones_cambio', ''))
        
        # ===== ALTERNATIVAS =====
        NativeDocumentGenerator._add_section_title(doc, "5. Alternativas Consideradas", level=1)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('alternativas_consideradas', ''))
        
        # ===== PLAN DE RETORNO =====
        NativeDocumentGenerator._add_section_title(doc, "6. Plan de Retorno", level=1)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('plan_retorno', ''))
        
        doc.add_page_break()
        
        # ===== RECURSOS =====
        NativeDocumentGenerator._add_section_title(doc, "7. Recursos", level=1)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('recursos', ''))
        
        # ===== PLAN DE IMPLEMENTACIÓN =====
        NativeDocumentGenerator._add_section_title(doc, "8. Plan de Implementación", level=1)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('plan_implementacion', ''))
        
        # ===== TIEMPO =====
        NativeDocumentGenerator._add_section_title(doc, "9. Tiempo de Duración del Cambio", level=1)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('tiempo_duracion', ''))
        
        doc.add_page_break()
        
        # ===== CHECKLIST 360° =====
        NativeDocumentGenerator._add_section_title(doc, "10. Checklist 360° - Análisis Integral del Cambio", level=1)
        checklist = data.get('checklist_360', [])
        if checklist:
            table = doc.add_table(rows=len(checklist)+1, cols=4)
            table.style = 'Table Grid'
            NativeDocumentGenerator._add_header_row(table, 0, ["N°", "Factor a Revisar", "Aplica", "Descripción del Impacto"])
            for i, item in enumerate(checklist):
                row = table.rows[i+1]
                row.cells[0].text = str(item.get('numero', i+1))
                row.cells[1].text = item.get('factor', '')
                row.cells[2].text = item.get('aplica', 'NO')
                row.cells[3].text = item.get('descripcion', '')
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)
        
        doc.add_paragraph()
        
        # ===== DOCUMENTOS IMPACTADOS =====
        NativeDocumentGenerator._add_section_title(doc, "11. Documentos Impactados por el Cambio", level=1)
        docs_imp = data.get('documentos_impactados', [])
        if docs_imp:
            table = doc.add_table(rows=len(docs_imp)+1, cols=4)
            table.style = 'Table Grid'
            NativeDocumentGenerator._add_header_row(table, 0, ["N°", "Documento", "Aplica", "Modificación Específica"])
            for i, item in enumerate(docs_imp):
                row = table.rows[i+1]
                row.cells[0].text = str(item.get('numero', i+1))
                row.cells[1].text = item.get('documento', '')
                row.cells[2].text = item.get('aplica', 'NO')
                row.cells[3].text = item.get('modificacion', '')
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)
        
        doc.add_page_break()
        
        # ===== RIESGOS DE CALIDAD =====
        NativeDocumentGenerator._add_section_title(doc, "12. Riesgos de Calidad y Controles", level=1)
        riesgos_cal = data.get('riesgos_controles', [])
        if riesgos_cal:
            table = doc.add_table(rows=len(riesgos_cal)+1, cols=3)
            table.style = 'Table Grid'
            NativeDocumentGenerator._add_header_row(table, 0, ["N°", "Riesgo Identificado", "Control / Plan de Acción"])
            for i, risk in enumerate(riesgos_cal):
                row = table.rows[i+1]
                row.cells[0].text = str(i+1)
                row.cells[1].text = risk.get('riesgo', '')
                row.cells[2].text = risk.get('control', '')
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)
        
        doc.add_paragraph()
        
        # ===== RIESGOS SHES =====
        NativeDocumentGenerator._add_section_title(doc, "13. Riesgos SHES y Medidas de Control", level=1)
        riesgos_shes = data.get('riesgos_shes', [])
        if riesgos_shes:
            table = doc.add_table(rows=len(riesgos_shes)+1, cols=4)
            table.style = 'Table Grid'
            NativeDocumentGenerator._add_header_row(table, 0, ["N°", "Riesgo Identificado", "Controles Propuestos", "Plazo"])
            for i, risk in enumerate(riesgos_shes):
                row = table.rows[i+1]
                row.cells[0].text = str(i+1)
                row.cells[1].text = risk.get('riesgo', '')
                row.cells[2].text = risk.get('control', '')
                row.cells[3].text = risk.get('plazo', '')
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)
        
        doc.add_paragraph()
        
        # ===== IMPACTO ESPERADO =====
        NativeDocumentGenerator._add_section_title(doc, "14. Impacto Esperado del Cambio", level=1)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('impacto_esperado', ''))
        
        # ===== RESUMEN EJECUTIVO =====
        NativeDocumentGenerator._add_section_title(doc, "15. Resumen Ejecutivo para Aprobación", level=1)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('resumen_ejecutivo', ''))
        
        # ===== IMÁGENES =====
        if images:
            doc.add_page_break()
            NativeDocumentGenerator._add_section_title(doc, "Anexo: Imágenes de Soporte", level=1)
            for img_info in images:
                try:
                    doc.add_picture(img_info['path'], width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption = doc.add_paragraph(img_info.get('desc', ''))
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        run.font.italic = True
                        run.font.size = Pt(10)
                    doc.add_paragraph()
                except Exception as e:
                    doc.add_paragraph(f"[Error al cargar imagen: {e}]")
        
        # Guardar en buffer
        output_buffer = BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        return output_buffer
    
    @staticmethod
    def generate_a3_docx(data, images=None):
        """Genera documento A3 en formato A4 Word nativo"""
        doc = Document()
        
        for section in doc.sections:
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # ===== PORTADA =====
        NativeDocumentGenerator._add_page_header(
            doc,
            "RESOLUCIÓN DE PROBLEMAS A3",
            "MEJORA CONTINUA"
        )
        
        # Información general
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        info_fields = [
            ("Título:", data.get('titulo', '')),
            ("Área:", data.get('area', '')),
            ("Autor:", data.get('autor', '')),
            ("Fecha:", data.get('fecha', '')),
        ]
        for i, (label, value) in enumerate(info_fields):
            cell_label = table.rows[i].cells[0]
            cell_value = table.rows[i].cells[1]
            cell_label.text = label
            cell_value.text = value
            NativeDocumentGenerator._set_cell_shading(cell_label, "E2E8F0")
            for p in cell_label.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            for p in cell_value.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # ===== SECCIONES A3 =====
        sections = [
            ("1. ANTECEDENTES", "antecedentes"),
            ("2. PROBLEMA ACTUAL", "problema_actual"),
            ("3. ANÁLISIS DE LA SITUACIÓN", "analisis_situacion"),
            ("4. OBJETIVOS", "objetivos"),
            ("5. ANÁLISIS DE CAUSA RAÍZ", "analisis_causa_raiz"),
            ("6. CONTRAMEDIDAS", "contramedidas"),
            ("7. RESULTADOS ESPERADOS", "resultados_esperados"),
            ("8. PLAN DE SEGUIMIENTO", "plan_seguimiento"),
            ("9. LECCIONES APRENDIDAS", "lecciones_aprendidas"),
            ("10. ESTANDARIZACIÓN", "estandarizacion"),
        ]
        
        for title, key in sections:
            NativeDocumentGenerator._add_section_title(doc, title, level=2)
            content = data.get(key, '')
            if content:
                NativeDocumentGenerator._add_content_paragraph(doc, content)
        
        # ===== IMÁGENES =====
        if images:
            doc.add_page_break()
            NativeDocumentGenerator._add_section_title(doc, "Anexo: Imágenes de Soporte", level=1)
            for img_info in images:
                try:
                    doc.add_picture(img_info['path'], width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption = doc.add_paragraph(img_info.get('desc', ''))
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        run.font.italic = True
                        run.font.size = Pt(10)
                    doc.add_paragraph()
                except Exception as e:
                    doc.add_paragraph(f"[Error al cargar imagen: {e}]")
        
        output_buffer = BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        return output_buffer
    
    @staticmethod
    def generate_kaizen_docx(data, images=None):
        """Genera documento Kaizen en formato A4 Word nativo"""
        doc = Document()
        
        for section in doc.sections:
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
        
        # ===== PORTADA =====
        NativeDocumentGenerator._add_page_header(
            doc,
            "SIMPLE KAIZEN",
            "Registro de Mejora Continua"
        )
        
        # Información general
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        info_fields = [
            ("Name (Título):", data.get('titulo', '')),
            ("Plant / Area:", data.get('area', '')),
            ("Date:", data.get('fecha', '')),
            ("Leader:", data.get('leader', '')),
            ("Team Members:", data.get('team_members', '')),
        ]
        for i, (label, value) in enumerate(info_fields):
            cell_label = table.rows[i].cells[0]
            cell_value = table.rows[i].cells[1]
            cell_label.text = label
            cell_value.text = value
            NativeDocumentGenerator._set_cell_shading(cell_label, "E2E8F0")
            for p in cell_label.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
            for p in cell_value.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # ===== OPORTUNIDAD =====
        NativeDocumentGenerator._add_section_title(doc, "Opportunity (Descripción del Problema)", level=2)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('descripcion_problema', ''))
        
        # ===== IMPROVEMENT =====
        NativeDocumentGenerator._add_section_title(doc, "Improvement (Solución Implementada)", level=2)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('solucion', ''))
        
        # ===== BENEFIT =====
        NativeDocumentGenerator._add_section_title(doc, "Benefit (Beneficios)", level=2)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('beneficios', ''))
        
        # ===== CLASIFICACIÓN =====
        doc.add_page_break()
        NativeDocumentGenerator._add_section_title(doc, "Clasificación del Kaizen", level=1)
        
        # Tabla de 8 Wastes
        NativeDocumentGenerator._add_section_title(doc, "8 Wastes (Desperdicios)", level=2)
        wastes = ["Motion", "Skills", "Inventory", "Transportation", 
                  "Over Production", "Over Processing", "Waiting", "Defects"]
        selected_waste = data.get('tipo_desperdicio', '').lower()
        
        table = doc.add_table(rows=len(wastes)+1, cols=2)
        table.style = 'Table Grid'
        NativeDocumentGenerator._add_header_row(table, 0, ["Desperdicio", "Seleccionado"])
        for i, waste in enumerate(wastes):
            row = table.rows[i+1]
            row.cells[0].text = waste
            is_selected = waste.lower() in selected_waste
            row.cells[1].text = "✓ X" if is_selected else ""
            if is_selected:
                NativeDocumentGenerator._set_cell_shading(row.cells[1], "DCFCE7")
                for p in row.cells[1].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = DocxRGBColor(0x16, 0x65, 0x34)
        
        doc.add_paragraph()
        
        # Tabla BTO
        NativeDocumentGenerator._add_section_title(doc, "BTO Impact", level=2)
        bto_categories = [
            "Safe and Sustainable",
            "People & Culture",
            "Network Optimisation",
            "Supply Chain and Manufacturing Excellence"
        ]
        selected_bto = data.get('impacto_bto', '').lower()
        
        table = doc.add_table(rows=len(bto_categories)+1, cols=2)
        table.style = 'Table Grid'
        NativeDocumentGenerator._add_header_row(table, 0, ["Categoría BTO", "Seleccionado"])
        for i, bto in enumerate(bto_categories):
            row = table.rows[i+1]
            row.cells[0].text = bto
            is_selected = bto.lower() in selected_bto
            row.cells[1].text = "✓ X" if is_selected else ""
            if is_selected:
                NativeDocumentGenerator._set_cell_shading(row.cells[1], "DCFCE7")
                for p in row.cells[1].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.bold = True
                        run.font.color.rgb = DocxRGBColor(0x16, 0x65, 0x34)
        
        doc.add_paragraph()
        
        # ===== PRÓXIMOS PASOS =====
        NativeDocumentGenerator._add_section_title(doc, "Próximos Pasos", level=2)
        NativeDocumentGenerator._add_content_paragraph(doc, data.get('proximos_pasos', ''))
        
        # ===== IMÁGENES =====
        if images:
            doc.add_page_break()
            NativeDocumentGenerator._add_section_title(doc, "Anexo: Imágenes Antes / Después", level=1)
            for img_info in images:
                try:
                    doc.add_picture(img_info['path'], width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption = doc.add_paragraph(img_info.get('desc', ''))
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in caption.runs:
                        run.font.italic = True
                        run.font.size = Pt(10)
                    doc.add_paragraph()
                except Exception as e:
                    doc.add_paragraph(f"[Error al cargar imagen: {e}]")
        
        output_buffer = BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        return output_buffer

# =============================================================================
# EXPORTADOR PDF NATIVO
# =============================================================================
class NativePDFExporter:
    """Genera PDFs nativos usando ReportLab"""
    
    @staticmethod
    def generate_moc_pdf(data, images=None):
        """Genera PDF de MoC nativo"""
        if not REPORTLAB_AVAILABLE:
            return None
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=50, leftMargin=50,
                                topMargin=50, bottomMargin=50)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle', parent=styles['Heading1'],
            fontSize=20, textColor=colors.HexColor('#1a5f7a'),
            spaceAfter=20, alignment=TA_CENTER, fontName='Helvetica-Bold'
        )
        subtitle_style = ParagraphStyle(
            'CustomSubtitle', parent=styles['Normal'],
            fontSize=11, textColor=colors.HexColor('#64748b'),
            spaceAfter=10, alignment=TA_CENTER
        )
        heading_style = ParagraphStyle(
            'CustomHeading', parent=styles['Heading2'],
            fontSize=13, textColor=colors.HexColor('#1a5f7a'),
            spaceAfter=8, spaceBefore=12, fontName='Helvetica-Bold'
        )
        body_style = ParagraphStyle(
            'CustomBody', parent=styles['BodyText'],
            fontSize=10, leading=13, alignment=TA_JUSTIFY, fontName='Helvetica'
        )
        
        story = []
        
        # Portada
        story.append(Paragraph("MANAGEMENT OF CHANGE (MoC)", title_style))
        story.append(Paragraph("Gestión del Cambio", subtitle_style))
        story.append(Spacer(1, 15))
        
        # Tabla de información
        info_data = [
            ["Título:", data.get('moc_title', '')],
            ["Número:", data.get('moc_number', '')],
            ["Fecha:", data.get('fecha', '')],
            ["Naturaleza:", data.get('naturaleza', '').upper()],
            ["Originador:", data.get('originador', '')],
        ]
        info_table = Table(info_data, colWidths=[120, 350])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#E2E8F0')),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 15))
        
        # Equipo
        story.append(Paragraph("Equipo de Revisión", heading_style))
        equipo_text = f"""
        Producción: {data.get('produccion', 'N/A')}<br/>
        Specialist SHES: {data.get('specialist_shes', 'N/A')}<br/>
        Mantenimiento: {data.get('mantenimiento', 'N/A')}<br/>
        Revisores: {data.get('revisores', 'N/A')}<br/>
        Experto Aprobador: {data.get('experto_aprobador', 'N/A')}
        """
        story.append(Paragraph(equipo_text, body_style))
        story.append(Spacer(1, 10))
        
        # Secciones
        sections = [
            ("1. Condición Actual", "condicion_actual"),
            ("2. Condición Propuesta", "condicion_propuesta"),
            ("3. Justificación de la MoC", "justificacion_moc"),
            ("4. Descripción del Problema", "descripcion_problema"),
            ("5. Razones del Cambio", "razones_cambio"),
            ("6. Alternativas Consideradas", "alternativas_consideradas"),
            ("7. Plan de Retorno", "plan_retorno"),
            ("8. Recursos", "recursos"),
            ("9. Plan de Implementación", "plan_implementacion"),
            ("10. Tiempo de Duración", "tiempo_duracion"),
            ("14. Impacto Esperado", "impacto_esperado"),
            ("15. Resumen Ejecutivo", "resumen_ejecutivo"),
        ]
        
        for title, key in sections:
            story.append(Paragraph(title, heading_style))
            content = data.get(key, '').replace('\n', '<br/>')
            story.append(Paragraph(content, body_style))
            story.append(Spacer(1, 8))
        
        # Checklist 360°
        story.append(Paragraph("11. Checklist 360° - Análisis Integral", heading_style))
        checklist = data.get('checklist_360', [])
        if checklist:
            chk_data = [["N°", "Factor", "Aplica", "Descripción"]]
            for item in checklist:
                chk_data.append([
                    str(item.get('numero', '')),
                    item.get('factor', ''),
                    item.get('aplica', 'NO'),
                    item.get('descripcion', '')
                ])
            chk_table = Table(chk_data, colWidths=[25, 180, 40, 225])
            chk_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a5f7a')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
            ]))
            story.append(chk_table)
            story.append(Spacer(1, 10))
        
        # Documentos impactados
        story.append(Paragraph("12. Documentos Impactados", heading_style))
        docs_imp = data.get('documentos_impactados', [])
        if docs_imp:
            doc_data = [["N°", "Documento", "Aplica", "Modificación"]]
            for item in docs_imp:
                doc_data.append([
                    str(item.get('numero', '')),
                    item.get('documento', ''),
                    item.get('aplica', 'NO'),
                    item.get('modificacion', '')
                ])
            doc_table = Table(doc_data, colWidths=[25, 150, 40, 255])
            doc_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a5f7a')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(doc_table)
            story.append(Spacer(1, 10))
        
        # Riesgos
        story.append(Paragraph("13. Riesgos de Calidad y SHES", heading_style))
        riesgos_cal = data.get('riesgos_controles', [])
        if riesgos_cal:
            risk_data = [["N°", "Riesgo", "Control"]]
            for i, risk in enumerate(riesgos_cal, 1):
                risk_data.append([str(i), risk.get('riesgo', ''), risk.get('control', '')])
            risk_table = Table(risk_data, colWidths=[25, 220, 225])
            risk_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a5f7a')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(risk_table)
        
        riesgos_shes = data.get('riesgos_shes', [])
        if riesgos_shes:
            story.append(Spacer(1, 10))
            shes_data = [["N°", "Riesgo SHES", "Control", "Plazo"]]
            for i, risk in enumerate(riesgos_shes, 1):
                shes_data.append([str(i), risk.get('riesgo', ''), risk.get('control', ''), risk.get('plazo', '')])
            shes_table = Table(shes_data, colWidths=[25, 160, 180, 105])
            shes_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a5f7a')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(shes_table)
        
        # Imágenes
        if images:
            story.append(PageBreak())
            story.append(Paragraph("Anexo: Imágenes de Soporte", heading_style))
            for img_info in images:
                try:
                    img = RLImage(img_info['path'], width=450)
                    story.append(img)
                    story.append(Paragraph(img_info.get('desc', ''), subtitle_style))
                    story.append(Spacer(1, 10))
                except Exception as e:
                    story.append(Paragraph(f"[Error: {e}]", body_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    @staticmethod
    def generate_a3_pdf(data, images=None):
        """Genera PDF de A3 nativo"""
        if not REPORTLAB_AVAILABLE:
            return None
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=50, leftMargin=50,
                                topMargin=50, bottomMargin=50)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle', parent=styles['Heading1'],
            fontSize=20, textColor=colors.HexColor('#1a5f7a'),
            spaceAfter=20, alignment=TA_CENTER, fontName='Helvetica-Bold'
        )
        subtitle_style = ParagraphStyle(
            'CustomSubtitle', parent=styles['Normal'],
            fontSize=11, textColor=colors.HexColor('#64748b'),
            spaceAfter=10, alignment=TA_CENTER
        )
        heading_style = ParagraphStyle(
            'CustomHeading', parent=styles['Heading2'],
            fontSize=13, textColor=colors.HexColor('#1a5f7a'),
            spaceAfter=8, spaceBefore=12, fontName='Helvetica-Bold'
        )
        body_style = ParagraphStyle(
            'CustomBody', parent=styles['BodyText'],
            fontSize=10, leading=13, alignment=TA_JUSTIFY, fontName='Helvetica'
        )
        
        story = []
        story.append(Paragraph("RESOLUCIÓN DE PROBLEMAS A3", title_style))
        story.append(Paragraph("MEJORA CONTINUA", subtitle_style))
        story.append(Spacer(1, 15))
        
        info_data = [
            ["Título:", data.get('titulo', '')],
            ["Área:", data.get('area', '')],
            ["Autor:", data.get('autor', '')],
            ["Fecha:", data.get('fecha', '')],
        ]
        info_table = Table(info_data, colWidths=[120, 350])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#E2E8F0')),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 15))
        
        sections = [
            ("1. ANTECEDENTES", "antecedentes"),
            ("2. PROBLEMA ACTUAL", "problema_actual"),
            ("3. ANÁLISIS DE LA SITUACIÓN", "analisis_situacion"),
            ("4. OBJETIVOS", "objetivos"),
            ("5. ANÁLISIS DE CAUSA RAÍZ", "analisis_causa_raiz"),
            ("6. CONTRAMEDIDAS", "contramedidas"),
            ("7. RESULTADOS ESPERADOS", "resultados_esperados"),
            ("8. PLAN DE SEGUIMIENTO", "plan_seguimiento"),
            ("9. LECCIONES APRENDIDAS", "lecciones_aprendidas"),
            ("10. ESTANDARIZACIÓN", "estandarizacion"),
        ]
        
        for title, key in sections:
            story.append(Paragraph(title, heading_style))
            content = data.get(key, '').replace('\n', '<br/>')
            story.append(Paragraph(content, body_style))
            story.append(Spacer(1, 8))
        
        if images:
            story.append(PageBreak())
            story.append(Paragraph("Anexo: Imágenes de Soporte", heading_style))
            for img_info in images:
                try:
                    img = RLImage(img_info['path'], width=450)
                    story.append(img)
                    story.append(Paragraph(img_info.get('desc', ''), subtitle_style))
                    story.append(Spacer(1, 10))
                except Exception as e:
                    story.append(Paragraph(f"[Error: {e}]", body_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    @staticmethod
    def generate_kaizen_pdf(data, images=None):
        """Genera PDF de Kaizen nativo"""
        if not REPORTLAB_AVAILABLE:
            return None
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                rightMargin=50, leftMargin=50,
                                topMargin=50, bottomMargin=50)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle', parent=styles['Heading1'],
            fontSize=20, textColor=colors.HexColor('#f59e0b'),
            spaceAfter=20, alignment=TA_CENTER, fontName='Helvetica-Bold'
        )
        subtitle_style = ParagraphStyle(
            'CustomSubtitle', parent=styles['Normal'],
            fontSize=11, textColor=colors.HexColor('#64748b'),
            spaceAfter=10, alignment=TA_CENTER
        )
        heading_style = ParagraphStyle(
            'CustomHeading', parent=styles['Heading2'],
            fontSize=13, textColor=colors.HexColor('#f59e0b'),
            spaceAfter=8, spaceBefore=12, fontName='Helvetica-Bold'
        )
        body_style = ParagraphStyle(
            'CustomBody', parent=styles['BodyText'],
            fontSize=10, leading=13, alignment=TA_JUSTIFY, fontName='Helvetica'
        )
        
        story = []
        story.append(Paragraph("SIMPLE KAIZEN", title_style))
        story.append(Paragraph("Registro de Mejora Continua", subtitle_style))
        story.append(Spacer(1, 15))
        
        info_data = [
            ["Name:", data.get('titulo', '')],
            ["Plant / Area:", data.get('area', '')],
            ["Date:", data.get('fecha', '')],
            ["Leader:", data.get('leader', '')],
            ["Team Members:", data.get('team_members', '')],
        ]
        info_table = Table(info_data, colWidths=[120, 350])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#FEF3C7')),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 15))
        
        story.append(Paragraph("Opportunity (Problema)", heading_style))
        story.append(Paragraph(data.get('descripcion_problema', '').replace('\n', '<br/>'), body_style))
        story.append(Spacer(1, 10))
        
        story.append(Paragraph("Improvement (Solución)", heading_style))
        story.append(Paragraph(data.get('solucion', '').replace('\n', '<br/>'), body_style))
        story.append(Spacer(1, 10))
        
        story.append(Paragraph("Benefit (Beneficios)", heading_style))
        story.append(Paragraph(data.get('beneficios', '').replace('\n', '<br/>'), body_style))
        story.append(Spacer(1, 10))
        
        story.append(Paragraph("Próximos Pasos", heading_style))
        story.append(Paragraph(data.get('proximos_pasos', '').replace('\n', '<br/>'), body_style))
        story.append(Spacer(1, 15))
        
        # Clasificación
        story.append(Paragraph("Clasificación del Kaizen", heading_style))
        
        wastes = ["Motion", "Skills", "Inventory", "Transportation",
                  "Over Production", "Over Processing", "Waiting", "Defects"]
        selected_waste = data.get('tipo_desperdicio', '').lower()
        waste_data = [["Desperdicio", "Seleccionado"]]
        for waste in wastes:
            is_selected = waste.lower() in selected_waste
            waste_data.append([waste, "✓ X" if is_selected else ""])
        waste_table = Table(waste_data, colWidths=[200, 100])
        waste_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f59e0b')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
        ]))
        story.append(waste_table)
        story.append(Spacer(1, 10))
        
        bto_categories = ["Safe and Sustainable", "People & Culture",
                          "Network Optimisation", "Supply Chain and Manufacturing Excellence"]
        selected_bto = data.get('impacto_bto', '').lower()
        bto_data = [["Categoría BTO", "Seleccionado"]]
        for bto in bto_categories:
            is_selected = bto.lower() in selected_bto
            bto_data.append([bto, "✓ X" if is_selected else ""])
        bto_table = Table(bto_data, colWidths=[250, 100])
        bto_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f59e0b')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
        ]))
        story.append(bto_table)
        
        if images:
            story.append(PageBreak())
            story.append(Paragraph("Anexo: Imágenes Antes / Después", heading_style))
            for img_info in images:
                try:
                    img = RLImage(img_info['path'], width=450)
                    story.append(img)
                    story.append(Paragraph(img_info.get('desc', ''), subtitle_style))
                    story.append(Spacer(1, 10))
                except Exception as e:
                    story.append(Paragraph(f"[Error: {e}]", body_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()

# =============================================================================
# SESSION STATE
# =============================================================================
def init_session_state():
    saved_config = LocalStorage.load_config()
    saved_history = LocalStorage.load_history()
    
    if saved_config and saved_config.get("gemini_model") in GeminiService.DEPRECATED_MODELS:
        saved_config["gemini_model"] = GeminiService.DEFAULT_MODEL
        LocalStorage.save_config(saved_config)
    
    defaults = {
        "page": "inicio",
        "config": saved_config or {
            "gemini_api_key": "",
            "gemini_model": GeminiService.DEFAULT_MODEL,
            "company_name": "",
            "department": "",
            "default_author": "",
            "default_area": "",
            "last_moc_number": 0,
            "last_a3_number": 0,
            "last_kaizen_number": 0,
            "spell_check": True,
            "auto_correct": True,
        },
        "history": saved_history or {"documents": []},
        "generated_data": {},
        "doc_meta": {},
        "doc_images": [],
        "doc_type": None,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value
    if saved_config is None:
        LocalStorage.save_config(st.session_state.config)
    if saved_history is None:
        LocalStorage.save_history(st.session_state.history)

init_session_state()

# =============================================================================
# INTERFAZ
# =============================================================================
def render_sidebar():
    config = st.session_state.config
    st.sidebar.markdown("""
<div style="text-align: center; padding: 1rem 0;">
<h2 style="color: #f8fafc; margin: 0; font-size: 1.3rem;">📋 GESTIÓN</h2>
<h2 style="color: #f8fafc; margin: 0; font-size: 1.3rem;">DOCUMENTAL</h2>
<p style="color: #94a3b8; margin-top: 0.5rem; font-size: 0.85rem;">MoC · A3 · Kaizen</p>
</div>
<hr style="border-color: #334155; margin: 1rem 0;">
""", unsafe_allow_html=True)
    st.sidebar.markdown("### 🧭 Navegación")
    nav_items = [
        ("🏠 Inicio", "inicio"),
        ("📋 Nueva MoC", "nueva_moc"),
        ("📊 Nueva Mejora A3", "nueva_a3"),
        ("⚡ Nuevo Kaizen", "nuevo_kaizen"),
        ("📁 Historial", "historial"),
        ("⚙️ Configuración", "configuracion"),
    ]
    for label, page_key in nav_items:
        if st.sidebar.button(label, key=f"nav_{page_key}", use_container_width=True):
            st.session_state.page = page_key
            st.rerun()
    model_name = GeminiService.MODELS.get(config.get("gemini_model", GeminiService.DEFAULT_MODEL), {}).get("name", "Gemini 2.5 Flash")
    st.sidebar.markdown(f"""
<div style="text-align: center; color: #64748b; font-size: 0.75rem;">
<p>Modelo IA: <span class="gemini-badge">{model_name}</span></p>
<p>v10.0.0 · Agosto 2026</p>
</div>
""", unsafe_allow_html=True)
    st.sidebar.markdown("""
<hr style="border-color: #334155; margin: 1rem 0;">
<div style="text-align: center; padding: 0.5rem;">
<p style="color: #64748b; font-size: 0.75rem; margin: 0;">
<strong style="color: #94a3b8;">CAVA</strong><br>
Especialistas en Robótica<br>y Automatización<br><br>
Diseñado por<br>
<strong style="color: #2e8bc0;">Roger Huamani</strong>
</p>
</div>
""", unsafe_allow_html=True)

def render_header():
    st.markdown("""
<div class="main-header">
<h1>🎯 Sistema de Gestión Documental</h1>
<p>Automatización inteligente de documentos MoC, Mejora A3 y Simple Kaizen con IA</p>
</div>
""", unsafe_allow_html=True)

def render_welcome():
    render_header()
    st.success("✅ **Versión 10.0.0**: Generación nativa de documentos sin depender de templates externos. Todos los documentos se generan en formato A4 estándar.")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("""
<div class="doc-card doc-card-moc">
<h3 style="color: #1a5f7a; margin-top: 0;">📋 Management of Change</h3>
<p style="color: #64748b; font-size: 0.9rem;">Documento completo con 15 secciones, Checklist 360° y análisis integral.</p>
<ul style="color: #475569; font-size: 0.85rem; padding-left: 1.2rem;">
<li>Formato A4 nativo</li><li>Checklist 360° (16 factores)</li><li>15 documentos impactados</li><li>Riesgos SHES detallados</li>
</ul>
</div>
""", unsafe_allow_html=True)
        if st.button("➕ Crear MoC", key="btn_moc", use_container_width=True):
            st.session_state.page = "nueva_moc"
            st.rerun()
    with col2:
        st.markdown("""
<div class="doc-card doc-card-a3">
<h3 style="color: #10b981; margin-top: 0;">📊 Mejora A3</h3>
<p style="color: #64748b; font-size: 0.9rem;">Formato A3 con análisis 5 Porqués y contramedidas SMART.</p>
<ul style="color: #475569; font-size: 0.85rem; padding-left: 1.2rem;">
<li>Formato A4 nativo</li><li>Análisis 5 Porqués</li><li>Contramedidas priorizadas</li><li>Plan de seguimiento</li>
</ul>
</div>
""", unsafe_allow_html=True)
        if st.button("➕ Crear A3", key="btn_a3", use_container_width=True):
            st.session_state.page = "nueva_a3"
            st.rerun()
    with col3:
        st.markdown("""
<div class="doc-card doc-card-kaizen">
<h3 style="color: #f59e0b; margin-top: 0;">⚡ Simple Kaizen</h3>
<p style="color: #64748b; font-size: 0.9rem;">Registro rápido con clasificación 8 Wastes y BTO.</p>
<ul style="color: #475569; font-size: 0.85rem; padding-left: 1.2rem;">
<li>Formato A4 nativo</li><li>8 Wastes</li><li>Impacto BTO</li><li>Beneficios medibles</li>
</ul>
</div>
""", unsafe_allow_html=True)
        if st.button("➕ Crear Kaizen", key="btn_kaizen", use_container_width=True):
            st.session_state.page = "nuevo_kaizen"
            st.rerun()
    st.markdown("<br>", unsafe_allow_html=True)
    docs = st.session_state.history.get("documents", [])
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("📋 MoC", len([d for d in docs if d.get("type") == "moc"]))
    with col2:
        st.metric("📊 A3", len([d for d in docs if d.get("type") == "a3"]))
    with col3:
        st.metric("⚡ Kaizen", len([d for d in docs if d.get("type") == "kaizen"]))
    with col4:
        st.metric("📁 Total", len(docs))

def auto_correct_text_input(label, value, key, height=100, help_text=""):
    config = st.session_state.config
    if config.get("auto_correct", True):
        st.markdown('<div class="auto-correct-badge">✨ Corrección ortográfica automática activa</div>', unsafe_allow_html=True)
    text = st.text_area(label, value=value, height=height, key=key, help=help_text)
    if config.get("auto_correct", True) and text.strip():
        corrected = Utils.correct_spelling_basic(text)
        if corrected != text:
            st.info(f"🔧 Texto corregido automáticamente")
            return corrected
    return text

def render_moc_form():
    config = st.session_state.config
    st.markdown('<div class="section-header"><h3>📋 Nueva Management of Change (MoC)</h3></div>', unsafe_allow_html=True)
    st.info("💡 Complete la información y describa el problema. La IA generará automáticamente los 15 campos del formato MDET en formato A4.")
    
    st.markdown("#### 1. Información General")
    col1, col2 = st.columns(2)
    with col1:
        moc_title = st.text_input("Título de la MoC:", placeholder="Ej: INSTALACIÓN DE INTERLOCKS DE SEGURIDAD")
        moc_number = st.text_input("Número:", value=Utils.generate_doc_number("moc"), disabled=True)
    with col2:
        naturaleza = st.selectbox("Naturaleza:", ["permanente", "temporal", "emergencia"])
        originador = st.text_input("Originador:", value=config.get("default_author", ""))
        fecha = st.text_input("Fecha:", value=Utils.format_date_short(), disabled=True)
    
    st.markdown("#### 2. Equipo de Revisión")
    col1, col2, col3 = st.columns(3)
    with col1:
        produccion = st.text_input("Producción:")
        specialist_shes = st.text_input("Specialist SHES:")
    with col2:
        mantenimiento = st.text_input("Mantenimiento:")
        revisores = st.text_input("Revisores:")
    with col3:
        experto_aprobador = st.text_input("Experto Aprobador:")
    
    st.markdown("#### 3. Contexto del Problema (DETALLADO)")
    st.warning("⚠️ **IMPORTANTE:** Describa el problema con el mayor detalle posible. La IA generará contenido basado EXCLUSIVAMENTE en su contexto.")
    problem_desc = auto_correct_text_input(
        "Contexto del problema identificado:",
        "",
        "moc_problem_desc",
        height=300,
        help_text="Describa: qué está pasando, equipos involucrados, riesgos, normas aplicables (ISO 13849), solución propuesta."
    )
    st.markdown("#### 4. Contexto Adicional (Opcional)")
    context = auto_correct_text_input("Información adicional:", "", "moc_context", height=120)
    st.markdown("#### 5. Alternativas Consideradas (Opcional)")
    alternativas = auto_correct_text_input("Alternativas ya evaluadas:", "", "moc_alternativas", height=100)
    
    st.markdown("#### 6. Imágenes de Soporte (Opcional)")
    uploaded_images = st.file_uploader("Seleccione imágenes:", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
    image_paths = []
    if uploaded_images:
        for idx, img_file in enumerate(uploaded_images, 1):
            img_path = f"/tmp/temp_moc_img_{moc_number}_{idx}.png"
            with open(img_path, "wb") as f:
                f.write(img_file.getbuffer())
            image_paths.append({"path": img_path, "desc": f"Figura {idx} - {img_file.name}"})
        st.success(f"📷 {len(image_paths)} imagen(es) cargada(s)")
    
    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("🤖 Generar Documento MoC con IA", type="primary", use_container_width=True):
        if not problem_desc.strip():
            st.error("❌ Describa el problema antes de generar.")
            return
        with st.spinner("🧠 El Ingeniero Senior IA está generando el documento MoC..."):
            gemini = GeminiService(config.get("gemini_api_key", ""), config.get("gemini_model"))
            equipo_data = {
                "produccion": produccion, "specialist_shes": specialist_shes,
                "mantenimiento": mantenimiento, "revisores": revisores,
                "experto_aprobador": experto_aprobador
            }
            result = gemini.generate_moc(problem_desc, context, json.dumps(equipo_data), alternativas)
            if result is None:
                st.error("❌ No se pudo generar el documento.")
                return
            st.session_state.generated_data = result
            st.session_state.doc_meta = {
                "moc_title": moc_title, "moc_number": moc_number, "naturaleza": naturaleza,
                "originador": originador, "fecha": fecha, **equipo_data
            }
            st.session_state.doc_images = image_paths
            st.session_state.doc_type = "moc"
            st.session_state.page = "revisar"
            st.rerun()

def render_a3_form():
    config = st.session_state.config
    st.markdown('<div class="section-header"><h3>📊 Nueva Mejora A3</h3></div>', unsafe_allow_html=True)
    st.info("💡 Describa el problema con detalle. La IA generará el documento A3 completo en formato A4.")
    
    col1, col2 = st.columns(2)
    with col1:
        a3_title = st.text_input("Título:")
        area = st.text_input("Área:", value=config.get("default_area", ""))
    with col2:
        autor = st.text_input("Autor:", value=config.get("default_author", ""))
        doc_number = st.text_input("Número:", value=Utils.generate_doc_number("a3"), disabled=True)
        fecha = st.text_input("Fecha:", value=Utils.format_date(), disabled=True)
    
    problem_desc = auto_correct_text_input("Describa el problema actual:", "", "a3_problem_desc", height=250)
    context = auto_correct_text_input("Contexto adicional:", "", "a3_context", height=100)
    
    st.markdown("#### Imágenes de Soporte (Opcional)")
    uploaded_images = st.file_uploader("Seleccione imágenes:", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
    image_paths = []
    if uploaded_images:
        for idx, img_file in enumerate(uploaded_images, 1):
            img_path = f"/tmp/temp_a3_img_{doc_number}_{idx}.png"
            with open(img_path, "wb") as f:
                f.write(img_file.getbuffer())
            image_paths.append({"path": img_path, "desc": f"Figura {idx} - {img_file.name}"})
        st.success(f"📷 {len(image_paths)} imagen(es) cargada(s)")
    
    if st.button("🤖 Generar Documento A3 con IA", type="primary", use_container_width=True):
        if not problem_desc.strip():
            st.error("❌ Describa el problema antes de generar.")
            return
        with st.spinner("🧠 Generando documento A3..."):
            gemini = GeminiService(config.get("gemini_api_key", ""), config.get("gemini_model"))
            result = gemini.generate_a3(problem_desc, context)
            if result is None:
                return
            st.session_state.generated_data = result
            st.session_state.doc_meta = {"titulo": a3_title, "area": area, "autor": autor, "doc_number": doc_number, "fecha": fecha}
            st.session_state.doc_images = image_paths
            st.session_state.doc_type = "a3"
            st.session_state.page = "revisar"
            st.rerun()

def render_kaizen_form():
    config = st.session_state.config
    st.markdown('<div class="section-header"><h3>⚡ Nuevo Simple Kaizen</h3></div>', unsafe_allow_html=True)
    st.info("💡 Describa la actividad de mejora realizada con detalle.")
    
    col1, col2 = st.columns(2)
    with col1:
        kaizen_title = st.text_input("Título (Name):")
        area = st.text_input("Plant/Area:", value=config.get("default_area", ""))
    with col2:
        leader = st.text_input("Leader:", value=config.get("default_author", ""))
        doc_number = st.text_input("Número:", value=Utils.generate_doc_number("kaizen"), disabled=True)
        fecha = st.text_input("Date:", value=Utils.format_date(), disabled=True)
        team_members = st.text_input("Team Members:")
    
    activity_desc = auto_correct_text_input("Describa la mejora realizada:", "", "kzn_activity_desc", height=250)
    
    col1, col2 = st.columns(2)
    with col1:
        tipo_desp = st.multiselect("Tipo de Desperdicio:",
                                   ["Motion", "Skills", "Inventory", "Transportation",
                                    "Over Production", "Over Processing", "Waiting", "Defects"])
    with col2:
        impacto_bto = st.selectbox("Impacto BTO:",
                                   ["Safe and Sustainable", "People & Culture",
                                    "Network Optimisation", "Supply Chain and Manufacturing Excellence"])
    
    beneficios = auto_correct_text_input("Beneficios:", "", "kzn_beneficios", height=120)
    
    st.markdown("#### Imágenes Antes/Después")
    col1, col2 = st.columns(2)
    with col1:
        img_antes = st.file_uploader("Imagen ANTES:", type=["png", "jpg", "jpeg"], key="img_antes")
    with col2:
        img_despues = st.file_uploader("Imagen DESPUÉS:", type=["png", "jpg", "jpeg"], key="img_despues")
    
    image_paths = []
    if img_antes:
        img_path = f"/tmp/temp_kzn_antes_{doc_number}.png"
        with open(img_path, "wb") as f:
            f.write(img_antes.getbuffer())
        image_paths.append({"path": img_path, "desc": "ANTES - Estado inicial"})
    if img_despues:
        img_path = f"/tmp/temp_kzn_despues_{doc_number}.png"
        with open(img_path, "wb") as f:
            f.write(img_despues.getbuffer())
        image_paths.append({"path": img_path, "desc": "DESPUÉS - Estado final"})
    
    if st.button("🤖 Generar Documento Kaizen con IA", type="primary", use_container_width=True):
        if not activity_desc.strip():
            st.error("❌ Describa la actividad antes de generar.")
            return
        with st.spinner("🧠 Generando documento Kaizen..."):
            gemini = GeminiService(config.get("gemini_api_key", ""), config.get("gemini_model"))
            result = gemini.generate_kaizen(activity_desc, "")
            if result is None:
                return
            result["tipo_desperdicio"] = ", ".join(tipo_desp) if tipo_desp else result.get("tipo_desperdicio", "")
            result["impacto_bto"] = impacto_bto
            result["leader"] = leader
            result["team_members"] = team_members
            result["beneficios"] = beneficios if beneficios else result.get("beneficios", "")
            st.session_state.generated_data = result
            st.session_state.doc_meta = {
                "titulo": kaizen_title, "area": area, "leader": leader,
                "team_members": team_members, "doc_number": doc_number, "fecha": fecha
            }
            st.session_state.doc_images = image_paths
            st.session_state.doc_type = "kaizen"
            st.session_state.page = "revisar"
            st.rerun()

def render_review():
    doc_type = st.session_state.doc_type
    data = st.session_state.get("generated_data", {})
    meta = st.session_state.get("doc_meta", {})
    images = st.session_state.get("doc_images", [])
    config = st.session_state.config
    type_names = {"moc": "MoC", "a3": "Mejora A3", "kaizen": "Simple Kaizen"}
    type_name = type_names.get(doc_type, "Documento")
    st.markdown(f'<div class="section-header"><h3>👁️ Revisar y Editar {type_name}</h3></div>', unsafe_allow_html=True)
    st.info("💡 Revise cada campo, realice correcciones manuales si es necesario y genere el documento final en DOCX o PDF.")
    
    if doc_type == "moc":
        _render_moc_review(data, meta, images, config)
    elif doc_type == "a3":
        _render_a3_review(data, meta, images, config)
    elif doc_type == "kaizen":
        _render_kaizen_review(data, meta, images, config)

def _spell_check_field(label, value, key_prefix, gemini):
    col1, col2 = st.columns([6, 1])
    with col1:
        text = st.text_area(label, value=value, height=120, key=f"{key_prefix}_field")
    with col2:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("✨ Corregir", key=f"{key_prefix}_spell"):
            with st.spinner("Corrigiendo..."):
                corrected = gemini.correct_spelling(text)
                st.session_state[f"{key_prefix}_corrected"] = corrected
                st.rerun()
    if st.session_state.get(f"{key_prefix}_corrected"):
        text = st.session_state[f"{key_prefix}_corrected"]
        del st.session_state[f"{key_prefix}_corrected"]
    return text

def _render_moc_review(data, meta, images, config):
    gemini = GeminiService(config.get("gemini_api_key", ""), config.get("gemini_model"))
    tabs = st.tabs(["📋 General", "📝 Contenido", "📊 Checklist", "📄 Documentos", "⚠️ Riesgos", "📷 Imágenes", "⚙️ Generar"])
    
    with tabs[0]:
        st.markdown("#### Información del Documento")
        meta["moc_title"] = st.text_input("Título:", value=meta.get("moc_title", ""))
        meta["moc_number"] = st.text_input("Número:", value=meta.get("moc_number", ""), disabled=True)
        meta["naturaleza"] = st.selectbox("Naturaleza:", ["permanente", "temporal", "emergencia"],
                                          index=["permanente", "temporal", "emergencia"].index(meta.get("naturaleza", "permanente")))
        meta["originador"] = st.text_input("Originador:", value=meta.get("originador", ""))
        st.markdown("#### Equipo de Revisión")
        meta["produccion"] = st.text_input("Producción:", value=meta.get("produccion", ""))
        meta["specialist_shes"] = st.text_input("Specialist SHES:", value=meta.get("specialist_shes", ""))
        meta["mantenimiento"] = st.text_input("Mantenimiento:", value=meta.get("mantenimiento", ""))
        meta["revisores"] = st.text_input("Revisores:", value=meta.get("revisores", ""))
        meta["experto_aprobador"] = st.text_input("Experto Aprobador:", value=meta.get("experto_aprobador", ""))
    
    with tabs[1]:
        fields = [
            ("Condición Actual", "condicion_actual"),
            ("Condición Propuesta", "condicion_propuesta"),
            ("Justificación de la MoC", "justificacion_moc"),
            ("Descripción del Problema", "descripcion_problema"),
            ("Razones del Cambio", "razones_cambio"),
            ("Alternativas Consideradas", "alternativas_consideradas"),
            ("Plan de Retorno", "plan_retorno"),
            ("Recursos", "recursos"),
            ("Plan de Implementación", "plan_implementacion"),
            ("Tiempo de Duración", "tiempo_duracion"),
            ("Impacto Esperado", "impacto_esperado"),
            ("Resumen Ejecutivo", "resumen_ejecutivo"),
        ]
        for label, key in fields:
            st.markdown(f"#### {label}")
            data[key] = _spell_check_field("", data.get(key, ""), f"moc_{key}", gemini)
    
    with tabs[2]:
        st.markdown("#### Checklist 360° - 16 Factores")
        checklist = data.get("checklist_360", [])
        updated_checklist = []
        for item in checklist:
            st.markdown(f"**{item.get('numero')}. {item.get('factor')}**")
            col1, col2 = st.columns([1, 3])
            with col1:
                aplica = st.selectbox("Aplica:", ["SI", "NO"],
                                      index=0 if item.get("aplica", "NO") == "SI" else 1,
                                      key=f"chk_{item.get('numero')}")
            with col2:
                desc = st.text_input("Descripción:", value=item.get("descripcion", ""),
                                     key=f"chk_desc_{item.get('numero')}")
            updated_checklist.append({"numero": item.get("numero"), "factor": item.get("factor"),
                                      "aplica": aplica, "descripcion": desc if aplica == "SI" else ""})
        data["checklist_360"] = updated_checklist
    
    with tabs[3]:
        st.markdown("#### Documentos Impactados - 15 Documentos")
        docs_imp = data.get("documentos_impactados", [])
        updated_docs = []
        for item in docs_imp:
            st.markdown(f"**{item.get('numero')}. {item.get('documento')}**")
            col1, col2 = st.columns([1, 3])
            with col1:
                aplica = st.selectbox("Aplica:", ["SI", "NO"],
                                      index=0 if item.get("aplica", "NO") == "SI" else 1,
                                      key=f"doc_{item.get('numero')}")
            with col2:
                modif = st.text_input("Modificación:", value=item.get("modificacion", ""),
                                      key=f"doc_mod_{item.get('numero')}")
            updated_docs.append({"numero": item.get("numero"), "documento": item.get("documento"),
                                 "aplica": aplica, "modificacion": modif if aplica == "SI" else ""})
        data["documentos_impactados"] = updated_docs
    
    with tabs[4]:
        st.markdown("#### Riesgos de Calidad")
        risks_cal = data.get("riesgos_controles", [])
        updated_cal = []
        for i, risk in enumerate(risks_cal):
            st.markdown(f"**Riesgo {i+1}**")
            col1, col2, col3 = st.columns([2, 2, 1])
            with col1:
                r_riesgo = st.text_input(f"Riesgo:", value=risk.get("riesgo", ""), key=f"rcal_r_{i}")
            with col2:
                r_control = st.text_input(f"Control:", value=risk.get("control", ""), key=f"rcal_c_{i}")
            with col3:
                r_plazo = st.text_input(f"Plazo:", value=risk.get("plazo", ""), key=f"rcal_p_{i}")
            updated_cal.append({"riesgo": r_riesgo, "control": r_control, "plazo": r_plazo})
        data["riesgos_controles"] = updated_cal
        
        st.markdown("#### Riesgos SHES")
        risks_shes = data.get("riesgos_shes", [])
        updated_shes = []
        for i, risk in enumerate(risks_shes):
            st.markdown(f"**Riesgo SHES {i+1}**")
            col1, col2, col3 = st.columns([2, 2, 1])
            with col1:
                s_riesgo = st.text_input(f"Riesgo:", value=risk.get("riesgo", ""), key=f"rshes_r_{i}")
            with col2:
                s_control = st.text_input(f"Control:", value=risk.get("control", ""), key=f"rshes_c_{i}")
            with col3:
                s_plazo = st.text_input(f"Plazo:", value=risk.get("plazo", ""), key=f"rshes_p_{i}")
            updated_shes.append({"riesgo": s_riesgo, "control": s_control, "plazo": s_plazo})
        data["riesgos_shes"] = updated_shes
    
    with tabs[5]:
        st.markdown("#### Imágenes Cargadas")
        if images:
            for idx, img_info in enumerate(images, 1):
                st.image(img_info["path"], caption=f"Figura {idx}: {img_info['desc']}", width=400)
        else:
            st.info("No se cargaron imágenes")
    
    with tabs[6]:
        st.markdown("#### Generar Documento Final (Formato A4)")
        st.success("✅ Documento listo para generar en formato A4 estándar")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            if st.button("📄 DOCX Español", type="primary", use_container_width=True):
                _finalize_document(data, meta, images, "es", "moc", "docx")
        with col2:
            if st.button("📄 PDF Español", type="primary", use_container_width=True):
                _finalize_document(data, meta, images, "es", "moc", "pdf")
        with col3:
            if st.button("🇺🇸 DOCX Inglés", type="secondary", use_container_width=True):
                _finalize_document(data, meta, images, "en", "moc", "docx")
        with col4:
            if st.button("🔄 Regenerar", use_container_width=True):
                st.session_state.page = "nueva_moc"
                st.rerun()
    
    st.session_state.generated_data = data
    st.session_state.doc_meta = meta

def _render_a3_review(data, meta, images, config):
    gemini = GeminiService(config.get("gemini_api_key", ""), config.get("gemini_model"))
    tabs = st.tabs(["📋 General", "📝 Contenido", "📷 Imágenes", "⚙️ Generar"])
    with tabs[0]:
        meta["titulo"] = st.text_input("Título:", value=meta.get("titulo", ""))
        meta["area"] = st.text_input("Área:", value=meta.get("area", ""))
        meta["autor"] = st.text_input("Autor:", value=meta.get("autor", ""))
    with tabs[1]:
        sections = [
            ("Antecedentes", "antecedentes"), ("Problema Actual", "problema_actual"),
            ("Análisis de Situación", "analisis_situacion"), ("Objetivos", "objetivos"),
            ("Análisis Causa Raíz", "analisis_causa_raiz"), ("Contramedidas", "contramedidas"),
            ("Resultados Esperados", "resultados_esperados"), ("Plan de Seguimiento", "plan_seguimiento"),
            ("Lecciones Aprendidas", "lecciones_aprendidas"), ("Estandarización", "estandarizacion"),
        ]
        for label, key in sections:
            st.markdown(f"**{label}**")
            data[key] = _spell_check_field("", data.get(key, ""), f"a3_{key}", gemini)
    with tabs[2]:
        if images:
            for idx, img_info in enumerate(images, 1):
                st.image(img_info["path"], caption=f"Figura {idx}: {img_info['desc']}", width=400)
        else:
            st.info("No se cargaron imágenes")
    with tabs[3]:
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            if st.button("📄 DOCX Español", type="primary", use_container_width=True):
                _finalize_document(data, meta, images, "es", "a3", "docx")
        with col2:
            if st.button("📄 PDF Español", type="primary", use_container_width=True):
                _finalize_document(data, meta, images, "es", "a3", "pdf")
        with col3:
            if st.button("🇺🇸 DOCX Inglés", type="secondary", use_container_width=True):
                _finalize_document(data, meta, images, "en", "a3", "docx")
        with col4:
            if st.button("🔄 Regenerar", use_container_width=True):
                st.session_state.page = "nueva_a3"
                st.rerun()
    st.session_state.generated_data = data
    st.session_state.doc_meta = meta

def _render_kaizen_review(data, meta, images, config):
    gemini = GeminiService(config.get("gemini_api_key", ""), config.get("gemini_model"))
    tabs = st.tabs(["📋 General", "📝 Contenido", "📷 Imágenes", "⚙️ Generar"])
    with tabs[0]:
        meta["titulo"] = st.text_input("Título:", value=meta.get("titulo", ""))
        meta["area"] = st.text_input("Plant/Area:", value=meta.get("area", ""))
        meta["leader"] = st.text_input("Leader:", value=meta.get("leader", ""))
        meta["team_members"] = st.text_input("Team Members:", value=meta.get("team_members", ""))
    with tabs[1]:
        st.markdown("**Descripción del Problema**")
        data["descripcion_problema"] = _spell_check_field("", data.get("descripcion_problema", ""), "kzn_desc", gemini)
        st.markdown("**Solución**")
        data["solucion"] = _spell_check_field("", data.get("solucion", ""), "kzn_sol", gemini)
        st.markdown("**Beneficios**")
        data["beneficios"] = _spell_check_field("", data.get("beneficios", ""), "kzn_ben", gemini)
        st.markdown("**Tipo de Desperdicio**")
        data["tipo_desperdicio"] = st.text_input("", value=data.get("tipo_desperdicio", ""), key="kzn_desp")
        st.markdown("**Impacto BTO**")
        data["impacto_bto"] = st.text_input("", value=data.get("impacto_bto", ""), key="kzn_bto")
        st.markdown("**Próximos Pasos**")
        data["proximos_pasos"] = _spell_check_field("", data.get("proximos_pasos", ""), "kzn_next", gemini)
    with tabs[2]:
        if images:
            for idx, img_info in enumerate(images, 1):
                st.image(img_info["path"], caption=f"{img_info['desc']}", width=400)
        else:
            st.info("No se cargaron imágenes")
    with tabs[3]:
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            if st.button("📄 DOCX Español", type="primary", use_container_width=True):
                _finalize_document(data, meta, images, "es", "kaizen", "docx")
        with col2:
            if st.button("📄 PDF Español", type="primary", use_container_width=True):
                _finalize_document(data, meta, images, "es", "kaizen", "pdf")
        with col3:
            if st.button("🇺🇸 DOCX Inglés", type="secondary", use_container_width=True):
                _finalize_document(data, meta, images, "en", "kaizen", "docx")
        with col4:
            if st.button("🔄 Regenerar", use_container_width=True):
                st.session_state.page = "nuevo_kaizen"
                st.rerun()
    st.session_state.generated_data = data
    st.session_state.doc_meta = meta

def _finalize_document(data, meta, images, language, doc_type, output_format):
    config = st.session_state.config
    gemini = GeminiService(config.get("gemini_api_key", ""), config.get("gemini_model"))
    
    with st.spinner(f"📄 Generando documento en formato A4..."):
        final_data = {**meta, **data}
        if language == "en":
            st.info("🌐 Traduciendo documento al inglés...")
            final_data = gemini.translate_document(final_data)
        
        generator = NativeDocumentGenerator()
        pdf_exporter = NativePDFExporter()
        
        buffer = None
        ext = output_format
        mime = ""
        
        if doc_type == "moc":
            if output_format == "docx":
                buffer = generator.generate_moc_docx(final_data, images)
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                pdf_bytes = pdf_exporter.generate_moc_pdf(final_data, images)
                if pdf_bytes:
                    buffer = BytesIO(pdf_bytes)
                    mime = "application/pdf"
        elif doc_type == "a3":
            if output_format == "docx":
                buffer = generator.generate_a3_docx(final_data, images)
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                pdf_bytes = pdf_exporter.generate_a3_pdf(final_data, images)
                if pdf_bytes:
                    buffer = BytesIO(pdf_bytes)
                    mime = "application/pdf"
        else:  # kaizen
            if output_format == "docx":
                buffer = generator.generate_kaizen_docx(final_data, images)
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                pdf_bytes = pdf_exporter.generate_kaizen_pdf(final_data, images)
                if pdf_bytes:
                    buffer = BytesIO(pdf_bytes)
                    mime = "application/pdf"
        
        if buffer is None:
            st.error("❌ No se pudo generar el documento.")
            return
        
        lang_suffix = "en" if language == "en" else "es"
        filename = f"{meta.get('moc_number', meta.get('doc_number', 'DOC'))}_{lang_suffix}.{ext}"
        
        doc_info = {
            "type": doc_type,
            "title": meta.get("moc_title", meta.get("titulo", "Sin título")),
            "number": meta.get("moc_number", meta.get("doc_number", "")),
            "language": language,
            "format": ext,
            "filename": filename
        }
        Utils.add_to_history(doc_info)
        st.success(f"✅ Documento generado exitosamente: {filename}")
        st.download_button(
            label=f"📥 Descargar {ext.upper()}",
            data=buffer,
            file_name=filename,
            mime=mime,
            use_container_width=True
        )

def render_history():
    st.markdown('<div class="section-header"><h3>📁 Historial de Documentos</h3></div>', unsafe_allow_html=True)
    docs = st.session_state.history.get("documents", [])
    
    export_data = {
        "config": st.session_state.config,
        "history": st.session_state.history,
        "export_date": datetime.now().isoformat(),
        "version": "10.0.0"
    }
    st.download_button(
        label="📥 Exportar backup (JSON)",
        data=json.dumps(export_data, indent=2, ensure_ascii=False),
        file_name=f"backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
        mime="application/json",
        use_container_width=True
    )
    
    st.markdown("#### Restaurar desde archivo")
    restore_file = st.file_uploader("Seleccione archivo de backup:", type=["json"], key="restore_file")
    if restore_file:
        try:
            restore_data = json.loads(restore_file.read())
            if "config" in restore_data:
                st.session_state.config = restore_data["config"]
                LocalStorage.save_config(restore_data["config"])
            if "history" in restore_data:
                st.session_state.history = restore_data["history"]
                LocalStorage.save_history(restore_data["history"])
            st.success("✅ Restauración completada.")
            st.rerun()
        except Exception as e:
            st.error(f"❌ Error: {e}")
    
    if not docs:
        st.info("📭 No hay documentos generados aún.")
        return
    
    for doc in docs:
        type_emoji = {"moc": "📋", "a3": "📊", "kaizen": "⚡"}.get(doc.get("type"), "📄")
        lang_flag = "🇪🇸" if doc.get("language") == "es" else "🇺🇸"
        st.markdown(f"""
<div class="history-item">
<h4 style="margin: 0;">{type_emoji} {doc.get('title', 'Sin título')}</h4>
<p style="margin: 0.25rem 0; color: #64748b; font-size: 0.9rem;">
{doc.get('number', '')} · {lang_flag} · {doc.get('format', '').upper()} · {doc.get('timestamp', '')[:10]}
</p>
</div>
""", unsafe_allow_html=True)
        if st.button("🗑️ Eliminar", key=f"del_{doc.get('id', 'x')}"):
            Utils.delete_from_history(doc.get('id'))
            st.rerun()

def render_settings():
    st.markdown('<div class="section-header"><h3>⚙️ Configuración</h3></div>', unsafe_allow_html=True)
    config = st.session_state.config
    tabs = st.tabs(["🔑 API Gemini", "🏢 Empresa", "💾 Backup"])
    
    with tabs[0]:
        st.markdown("#### API Key Gemini")
        st.info("💡 Obtenga su API Key en [Google AI Studio](https://aistudio.google.com/)")
        api_key = st.text_input("API Key:", value=config.get("gemini_api_key", ""), type="password")
        
        if st.button("🔌 Probar Conexión API", use_container_width=True):
            if not api_key:
                st.error("⚠️ Ingrese una API Key primero.")
            else:
                with st.spinner("Probando..."):
                    try:
                        import requests
                        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={api_key}"
                        payload = {"contents": [{"parts": [{"text": "Responde solo 'OK'"}]}]}
                        resp = requests.post(url, json=payload, timeout=15)
                        if resp.status_code == 200:
                            st.success("✅ ¡Conexión exitosa!")
                        else:
                            st.error(f"❌ Error {resp.status_code}: {resp.text}")
                    except Exception as e:
                        st.error(f"❌ Error: {e}")
        
        st.markdown("#### Selección de Modelo")
        current_model = config.get("gemini_model", GeminiService.DEFAULT_MODEL)
        col1, col2 = st.columns(2)
        for i, (model_id, model_info) in enumerate(GeminiService.MODELS.items()):
            is_selected = current_model == model_id
            with [col1, col2][i]:
                st.markdown(f"**{model_info['name']}** - {model_info['desc']}")
                if st.button(f"{'✓ ' if is_selected else ''}Seleccionar", key=f"sel_{model_id}", use_container_width=True):
                    config["gemini_model"] = model_id
                    st.session_state.config = config
                    LocalStorage.save_config(config)
                    st.success(f"✅ Modelo: {model_info['name']}")
                    st.rerun()
        
        if st.button("💾 Guardar API Key", type="primary", use_container_width=True):
            config["gemini_api_key"] = api_key
            st.session_state.config = config
            LocalStorage.save_config(config)
            st.success("✅ API Key guardada")
    
    with tabs[1]:
        st.markdown("#### Datos de la Empresa")
        company = st.text_input("Empresa:", value=config.get("company_name", ""))
        dept = st.text_input("Departamento:", value=config.get("department", ""))
        author = st.text_input("Autor por defecto:", value=config.get("default_author", ""))
        area = st.text_input("Área por defecto:", value=config.get("default_area", ""))
        if st.button("💾 Guardar Datos", type="primary", use_container_width=True):
            config["company_name"] = company
            config["department"] = dept
            config["default_author"] = author
            config["default_area"] = area
            st.session_state.config = config
            LocalStorage.save_config(config)
            st.success("✅ Datos guardados")
    
    with tabs[2]:
        st.markdown("#### 🗑️ Zona de Peligro")
        st.warning("⚠️ Las siguientes acciones son irreversibles.")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🗑️ Borrar Historial", use_container_width=True):
                st.session_state.history = {"documents": []}
                LocalStorage.save_history({"documents": []})
                st.success("✅ Historial borrado")
                st.rerun()
        with col2:
            if st.button("🗑️ Borrar Configuración", use_container_width=True):
                st.session_state.config = {
                    "gemini_api_key": "",
                    "gemini_model": GeminiService.DEFAULT_MODEL,
                    "company_name": "",
                    "department": "",
                    "default_author": "",
                    "default_area": "",
                    "last_moc_number": 0,
                    "last_a3_number": 0,
                    "last_kaizen_number": 0,
                    "spell_check": True,
                    "auto_correct": True,
                }
                LocalStorage.save_config(st.session_state.config)
                st.success("✅ Configuración restaurada")
                st.rerun()

def main():
    render_sidebar()
    page = st.session_state.page
    if page == "inicio":
        render_welcome()
    elif page == "nueva_moc":
        render_moc_form()
    elif page == "nueva_a3":
        render_a3_form()
    elif page == "nuevo_kaizen":
        render_kaizen_form()
    elif page == "revisar":
        render_review()
    elif page == "historial":
        render_history()
    elif page == "configuracion":
        render_settings()
    else:
        render_welcome()
    st.markdown("""
<div class="app-footer">
<p><strong>CAVA</strong> - Especialistas en Robótica y Automatización</p>
<p>Diseñado por <strong>Roger Huamani</strong> | v10.0.0</p>
<p style="font-size: 0.75rem; color: #94a3b8;">
Generación nativa de documentos en formato A4 sin templates externos.<br>
Soporte para español e inglés. Exportación a DOCX y PDF.
</p>
</div>
""", unsafe_allow_html=True)

if __name__ == "__main__":
    main()
